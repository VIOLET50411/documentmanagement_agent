from pathlib import Path

import pytest

from app.ingestion.parsers.docx_parser import DocxParser
from app.ingestion.parsers.excel_parser import ExcelParser
from app.ingestion.parsers.ocr_parser import OCRParser
from app.ingestion.parsers.pdf_parser import PDFParser


def test_pdf_parser_fallback_notice_is_clean_text(tmp_path: Path):
    parser = PDFParser()
    notice = parser._fallback_notice(tmp_path / "sample.pdf", reason="pdf_text_unavailable")
    assert notice[0]["type"] == "ocr_notice"
    assert "待 OCR 或人工复核" in notice[0]["text"]


def test_ocr_parser_fallback_notice_is_clean_text(tmp_path: Path):
    parser = OCRParser()
    notice = parser._fallback_notice(tmp_path / "scan.pdf", reason="ocr_engine_unavailable")
    assert notice["type"] == "ocr_notice"
    assert "需要 OCR 才能提取正文" in notice["text"]


def test_ocr_parser_converts_pdf_pages_before_calling_engine(tmp_path: Path, monkeypatch):
    parser = OCRParser()
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 mock")

    class DummyPage:
        def save(self, path, format="PNG"):
            Path(path).write_bytes(b"png")

    class DummyEngine:
        def __init__(self):
            self.paths = []

        def ocr(self, path, cls=True):
            self.paths.append(Path(path).name)
            return [[("box", ("识别文本", 0.98))]]

    engine = DummyEngine()
    monkeypatch.setattr("app.ingestion.parsers.ocr_parser.convert_from_path", lambda _path: [DummyPage(), DummyPage()])

    result = parser._ocr_path(engine, pdf_path)

    assert len(result) == 2
    assert engine.paths == ["page-1.png", "page-2.png"]


def test_pdf_parser_falls_back_to_pypdf_when_unstructured_backend_errors(tmp_path: Path, monkeypatch):
    parser = PDFParser()

    monkeypatch.setattr(
        parser,
        "_parse_with_unstructured",
        lambda _path: (_ for _ in ()).throw(RuntimeError("pdfinfo missing")),
    )
    monkeypatch.setattr(
        parser,
        "_parse_with_pypdf",
        lambda path: [
            {
                "type": "paragraph",
                "text": "Hello PDF fallback",
                "metadata": {"page_number": 1, "section_title": path.stem, "parser": "pypdf"},
            }
        ],
    )

    rows = parser.parse(str(tmp_path / "sample.pdf"))

    assert rows
    assert rows[0]["metadata"]["parser"] == "pypdf"
    assert "Hello PDF fallback" in rows[0]["text"]


def test_pdf_parser_extracts_text_with_pypdf_on_minimal_pdf(tmp_path: Path):
    pytest.importorskip("pypdf")
    parser = PDFParser()
    pdf_path = tmp_path / "minimal.pdf"
    pdf_path.write_bytes(
        b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 41 >>
stream
BT
/F1 24 Tf
72 72 Td
(Hello PDF) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000241 00000 n 
0000000332 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
402
%%EOF
"""
    )

    rows = parser._parse_with_pypdf(pdf_path)

    assert rows
    assert any("Hello PDF" in row["text"] for row in rows)
    assert all(row["metadata"]["parser"] == "pypdf" for row in rows)


def test_excel_parser_import_fallback_is_clean_text(tmp_path: Path, monkeypatch):
    parser = ExcelParser()

    real_import = __import__

    def fake_import(name, *args, **kwargs):
        if name == "pandas":
            raise ImportError("missing pandas")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr("builtins.__import__", fake_import)
    rows = parser._parse_xlsx(tmp_path / "sheet.xlsx")
    assert rows[0]["type"] == "table"
    assert "需要 pandas/openpyxl" in rows[0]["text"]


def test_excel_parser_reads_xlsx_with_openpyxl(tmp_path: Path):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "budget.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "预算"
    sheet.append(["项目", "金额", "备注"])
    sheet.append(["差旅", 1200, "审批通过"])
    workbook.save(path)

    rows = ExcelParser().parse(str(path))

    assert rows
    assert rows[0]["metadata"]["parser"] == "openpyxl"
    assert rows[0]["metadata"]["sheet"] == "预算"
    assert "差旅" in rows[0]["text"]
    assert "1200" in rows[0]["text"]


def test_docx_parser_preserves_block_order_for_paragraphs_and_tables(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("总则", level=1)
    document.add_paragraph("第一段说明")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "报销"
    table.cell(1, 1).text = "需要审批"
    document.add_paragraph("第二段说明")
    document.save(path)

    rows = DocxParser().parse(str(path))

    assert [item["type"] for item in rows] == ["heading", "paragraph", "table", "paragraph"]
    assert rows[0]["text"] == "总则"
    assert rows[1]["text"] == "第一段说明"
    assert "字段" in rows[2]["text"]
    assert rows[3]["text"] == "第二段说明"
    assert [item["metadata"]["block_index"] for item in rows] == sorted(item["metadata"]["block_index"] for item in rows)
